#!/usr/bin/env python3
"""
fill_report.py — Deterministic Word report generator for FinOps Advanced Plugin

Replaces {{FIELD_NAME}} placeholders in a .docx template with values from a JSON
data file. Processes every location in the document: paragraphs, tables, headers,
footers, and text boxes. Output is deterministic: same inputs always produce the
same output.

Usage:
    python fill_report.py --template path/to/template.docx \
                          --data     path/to/data.json \
                          --output   path/to/output.docx

Exit codes:
    0  Success
    1  Missing required argument or file not found
    2  JSON parse error
    3  Template processing error

Dependencies:
    pip install python-docx
"""

import argparse
import json
import re
import sys
from copy import deepcopy
from pathlib import Path

# ---------- Dependency check ----------
try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print(
        "ERROR: python-docx is not installed.\n"
        "Run: pip install python-docx",
        file=sys.stderr,
    )
    sys.exit(1)


PLACEHOLDER_RE = re.compile(r"\{\{([A-Z0-9_]+)\}\}")


# ---------- Core replacement logic ----------

def replace_in_run(run, fields: dict[str, str]) -> list[str]:
    """Replace all placeholders in a single Run. Returns list of replaced keys."""
    replaced = []

    def _sub(m):
        key = m.group(1)
        if key in fields:
            replaced.append(key)
            return fields[key]
        return m.group(0)  # leave unknown placeholders intact

    run.text = PLACEHOLDER_RE.sub(_sub, run.text)
    return replaced


def replace_in_paragraph(para, fields: dict[str, str]) -> list[str]:
    """
    Replace placeholders across runs in a paragraph.

    Word sometimes splits a single placeholder across multiple runs
    (e.g. {{CLIENT}} becomes run "{{", run "CLIENT", run "}}").
    This function first merges the full paragraph text, checks for
    placeholders, then rebuilds with the replacement in the first run
    and empties the rest — preserving the original formatting of run 0.
    """
    replaced = []

    full_text = "".join(r.text for r in para.runs)
    if "{{" not in full_text:
        return replaced

    def _sub(m):
        key = m.group(1)
        if key in fields:
            replaced.append(key)
            return fields[key]
        return m.group(0)

    new_text = PLACEHOLDER_RE.sub(_sub, full_text)

    if new_text == full_text:
        return replaced  # nothing changed

    # Put the full replaced text into run 0, blank out all others
    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ""

    return replaced


def replace_in_table(table, fields: dict[str, str]) -> list[str]:
    """Walk every cell in a table (including nested tables)."""
    replaced = []
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                replaced.extend(replace_in_paragraph(para, fields))
            for nested_table in cell.tables:
                replaced.extend(replace_in_table(nested_table, fields))
    return replaced


def replace_in_section_headers_footers(section, fields: dict[str, str]) -> list[str]:
    """Replace placeholders in all headers and footers of a section."""
    replaced = []
    parts = [
        section.header, section.footer,
        section.first_page_header, section.first_page_footer,
        section.even_page_header, section.even_page_footer,
    ]
    for part in parts:
        if part is None:
            continue
        for para in part.paragraphs:
            replaced.extend(replace_in_paragraph(para, fields))
        for table in part.tables:
            replaced.extend(replace_in_table(table, fields))
    return replaced


def replace_in_text_boxes(doc, fields: dict[str, str]) -> list[str]:
    """Replace placeholders inside drawing/text-box elements in the document body."""
    replaced = []
    for txbx in doc.element.body.iter(qn("w:txbxContent")):
        for para_elem in txbx.iter(qn("w:p")):
            t_elements = list(para_elem.iter(qn("w:t")))
            if not t_elements:
                continue
            full_text = "".join(t.text or "" for t in t_elements)
            if "{{" not in full_text:
                continue

            def _sub(m, _fields=fields, _replaced=replaced):
                key = m.group(1)
                if key in _fields:
                    _replaced.append(key)
                    return _fields[key]
                return m.group(0)

            new_text = PLACEHOLDER_RE.sub(_sub, full_text)
            if new_text != full_text:
                t_elements[0].text = new_text
                for t in t_elements[1:]:
                    t.text = ""
    return replaced


# ---------- Main ----------

def load_fields(data_path: Path) -> dict[str, str]:
    """
    Load JSON data file and flatten to a dict of str→str.
    Nested keys are flattened with underscore: {"CLIENT": {"NAME": "Jane"}}
    becomes {"CLIENT_NAME": "Jane"}.
    All values are coerced to str. Keys are uppercased for consistent matching.
    Fields are returned sorted (determinism in logging).
    """
    raw = json.loads(data_path.read_text(encoding="utf-8"))

    flat: dict[str, str] = {}

    def _flatten(obj, prefix=""):
        if isinstance(obj, dict):
            # Sort keys for determinism
            for k in sorted(obj.keys()):
                _flatten(obj[k], f"{prefix}{k.upper()}_" if prefix else f"{k.upper()}_")
        else:
            key = prefix.rstrip("_")
            flat[key] = str(obj)

    _flatten(raw)
    return flat


def process(template_path: Path, data_path: Path, output_path: Path) -> None:
    print(f"Template : {template_path}", file=sys.stderr)
    print(f"Data     : {data_path}", file=sys.stderr)
    print(f"Output   : {output_path}", file=sys.stderr)

    fields = load_fields(data_path)
    print(f"Fields   : {len(fields)} loaded", file=sys.stderr)
    for key in sorted(fields):
        print(f"  {{{{ {key} }}}} = {fields[key]!r}", file=sys.stderr)

    doc = Document(str(template_path))
    all_replaced: list[str] = []

    # 1. Body paragraphs
    for para in doc.paragraphs:
        all_replaced.extend(replace_in_paragraph(para, fields))

    # 2. Body tables (sorted by position — deterministic)
    for table in doc.tables:
        all_replaced.extend(replace_in_table(table, fields))

    # 3. Headers and footers
    for section in doc.sections:
        all_replaced.extend(replace_in_section_headers_footers(section, fields))

    # 4. Text boxes
    all_replaced.extend(replace_in_text_boxes(doc, fields))

    # Report what was replaced (sorted for determinism)
    replaced_unique = sorted(set(all_replaced))
    print(f"\nReplaced : {len(replaced_unique)} unique placeholder(s)", file=sys.stderr)
    for key in replaced_unique:
        print(f"  {{{{ {key} }}}}", file=sys.stderr)

    # Report unfilled placeholders
    all_keys_in_template = set()
    for para in doc.paragraphs:
        for m in PLACEHOLDER_RE.finditer("".join(r.text for r in para.runs)):
            all_keys_in_template.add(m.group(1))

    unfilled = sorted(all_keys_in_template - set(replaced_unique))
    if unfilled:
        print(f"\nWARNING  : {len(unfilled)} placeholder(s) left unfilled:", file=sys.stderr)
        for key in unfilled:
            print(f"  {{{{ {key} }}}}", file=sys.stderr)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"\nSaved    : {output_path}", file=sys.stderr)


def parse_args(argv=None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Fill a .docx template with {{FIELD}} placeholders from a JSON data file.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument("--template", required=True, help="Path to the .docx template file")
    parser.add_argument("--data",     required=True, help="Path to the JSON data file")
    parser.add_argument("--output",   required=True, help="Path to write the filled .docx")
    return parser.parse_args(argv)


def main(argv=None) -> int:
    args = parse_args(argv)

    template_path = Path(args.template)
    data_path     = Path(args.data)
    output_path   = Path(args.output)

    if not template_path.exists():
        print(f"ERROR: Template not found: {template_path}", file=sys.stderr)
        return 1
    if not data_path.exists():
        print(f"ERROR: Data file not found: {data_path}", file=sys.stderr)
        return 1

    try:
        data_path.read_text(encoding="utf-8")
    except Exception as e:
        print(f"ERROR: Cannot read data file: {e}", file=sys.stderr)
        return 1

    try:
        json.loads(data_path.read_text(encoding="utf-8"))
    except json.JSONDecodeError as e:
        print(f"ERROR: Invalid JSON in data file: {e}", file=sys.stderr)
        return 2

    try:
        process(template_path, data_path, output_path)
    except Exception as e:
        print(f"ERROR: Template processing failed: {e}", file=sys.stderr)
        return 3

    return 0


if __name__ == "__main__":
    sys.exit(main())
